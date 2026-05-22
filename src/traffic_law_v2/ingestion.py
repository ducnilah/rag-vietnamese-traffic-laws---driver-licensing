from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, List, NamedTuple

from traffic_law_v2.models import Chunk, LegalMetadata, SourceDocument
from traffic_law_v2.text import chunk_by_tokens, normalize_text, stable_id, tokenize


CHAPTER_RE = re.compile(r"(?im)^\s*(Chương\s+[IVXLCDM\d]+[^\n]*)")
SECTION_RE = re.compile(r"(?im)^\s*(Mục\s+\d+[^\n]*)")
ARTICLE_RE = re.compile(r"(?im)^\s*(Điều\s+\d+[a-zA-Z]?\b[^\n]*)")
CLAUSE_RE = re.compile(r"(?m)^\s*(\d+)\.\s+")
POINT_RE = re.compile(r"(?m)^\s*([a-zđ])\)\s+")


class ArticleUnit(NamedTuple):
    chapter: str | None
    section: str | None
    article: str
    content: str


def load_documents(raw_dir: Path) -> List[SourceDocument]:
    docs: List[SourceDocument] = []
    for path in sorted(raw_dir.rglob("*")):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.suffix.lower() in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8")
        elif path.suffix.lower() == ".docx":
            text = _read_docx(path)
        else:
            continue
        text = normalize_text(text)
        if not text:
            continue
        source_id = stable_id(str(path), text[:500])
        metadata = LegalMetadata(
            source_id=source_id,
            source_path=str(path),
            document_title=_guess_title(text, path),
            legal_type=_guess_legal_type(text, path),
        )
        docs.append(SourceDocument(doc_id=source_id, text=text, metadata=metadata))
    return docs


def chunk_documents(
    docs: Iterable[SourceDocument],
    article_max_tokens: int = 1000,
    fallback_overlap: int = 50,
) -> List[Chunk]:
    chunks: List[Chunk] = []
    for doc in docs:
        chunk_index = 0
        article_units = _split_articles_with_hierarchy(doc.text)

        if not article_units:
            # Last-resort path for non-legal or malformed docs.
            for piece in chunk_by_tokens(doc.text, target_tokens=900, overlap=fallback_overlap):
                meta = doc.metadata.model_copy()
                chunks.append(
                    Chunk(
                        chunk_id=stable_id(doc.doc_id, str(chunk_index), piece[:120]),
                        doc_id=doc.doc_id,
                        text=piece,
                        metadata=meta,
                        chunk_index=chunk_index,
                        token_count=len(tokenize(piece)),
                    )
                )
                chunk_index += 1
            continue

        for unit in article_units:
            base_meta = doc.metadata.model_copy(
                update={
                    "chapter": unit.chapter,
                    "section": unit.section,
                    "article": unit.article,
                    "clause": None,
                    "point": None,
                }
            )
            article_tokens = len(tokenize(unit.content))

            # Priority 1: 1 chunk = 1 Điều
            if article_tokens <= article_max_tokens:
                chunks.append(
                    Chunk(
                        chunk_id=stable_id(doc.doc_id, str(chunk_index), unit.content[:120]),
                        doc_id=doc.doc_id,
                        text=unit.content,
                        metadata=base_meta,
                        chunk_index=chunk_index,
                        token_count=article_tokens,
                    )
                )
                chunk_index += 1
                continue

            # Priority 2: split by Khoản if Điều too long.
            clauses = _split_clauses(unit.content)
            if clauses:
                for clause_no, clause_text in clauses:
                    clause_meta = base_meta.model_copy(update={"clause": clause_no, "point": None})
                    clause_tokens = len(tokenize(clause_text))
                    if clause_tokens <= article_max_tokens:
                        chunks.append(
                            Chunk(
                                chunk_id=stable_id(doc.doc_id, str(chunk_index), clause_text[:120]),
                                doc_id=doc.doc_id,
                                text=clause_text,
                                metadata=clause_meta,
                                chunk_index=chunk_index,
                                token_count=clause_tokens,
                            )
                        )
                        chunk_index += 1
                        continue

                    # Priority 3: split by Điểm if Khoản still too long.
                    points = _split_points(clause_text)
                    if points:
                        for point_label, point_text in points:
                            point_meta = clause_meta.model_copy(update={"point": point_label})
                            point_tokens = len(tokenize(point_text))
                            if point_tokens <= article_max_tokens:
                                chunks.append(
                                    Chunk(
                                        chunk_id=stable_id(doc.doc_id, str(chunk_index), point_text[:120]),
                                        doc_id=doc.doc_id,
                                        text=point_text,
                                        metadata=point_meta,
                                        chunk_index=chunk_index,
                                        token_count=point_tokens,
                                    )
                                )
                                chunk_index += 1
                                continue

                            # Minimal-overlap fallback only when absolutely necessary.
                            for piece in chunk_by_tokens(point_text, target_tokens=900, overlap=fallback_overlap):
                                chunks.append(
                                    Chunk(
                                        chunk_id=stable_id(doc.doc_id, str(chunk_index), piece[:120]),
                                        doc_id=doc.doc_id,
                                        text=piece,
                                        metadata=point_meta,
                                        chunk_index=chunk_index,
                                        token_count=len(tokenize(piece)),
                                    )
                                )
                                chunk_index += 1
                    else:
                        # No point marker found; minimal-overlap fallback at clause level.
                        for piece in chunk_by_tokens(clause_text, target_tokens=900, overlap=fallback_overlap):
                            chunks.append(
                                Chunk(
                                    chunk_id=stable_id(doc.doc_id, str(chunk_index), piece[:120]),
                                    doc_id=doc.doc_id,
                                    text=piece,
                                    metadata=clause_meta,
                                    chunk_index=chunk_index,
                                    token_count=len(tokenize(piece)),
                                )
                            )
                            chunk_index += 1
            else:
                # No clause marker found; minimal-overlap fallback at article level.
                for piece in chunk_by_tokens(unit.content, target_tokens=900, overlap=fallback_overlap):
                    chunks.append(
                        Chunk(
                            chunk_id=stable_id(doc.doc_id, str(chunk_index), piece[:120]),
                            doc_id=doc.doc_id,
                            text=piece,
                            metadata=base_meta,
                            chunk_index=chunk_index,
                            token_count=len(tokenize(piece)),
                        )
                    )
                    chunk_index += 1
    return chunks


def _split_articles_with_hierarchy(text: str) -> List[ArticleUnit]:
    units: List[ArticleUnit] = []
    chapter: str | None = None
    section: str | None = None
    article_matches = list(ARTICLE_RE.finditer(text))
    if not article_matches:
        return units

    for idx, article_match in enumerate(article_matches):
        start = article_match.start()
        end = article_matches[idx + 1].start() if idx + 1 < len(article_matches) else len(text)
        content = text[start:end].strip()

        # Update hierarchy from region before this article.
        prev_start = article_matches[idx - 1].start() if idx > 0 else 0
        region_before = text[prev_start:start]
        chapter_candidate = _last_match(CHAPTER_RE, region_before)
        section_candidate = _last_match(SECTION_RE, region_before)
        if chapter_candidate:
            chapter = chapter_candidate
        if section_candidate:
            section = section_candidate

        article_header = normalize_text(article_match.group(1))
        units.append(ArticleUnit(chapter=chapter, section=section, article=article_header, content=content))
    return units


def _split_clauses(article_text: str) -> List[tuple[str, str]]:
    article_header = _first_line(article_text)
    matches = list(CLAUSE_RE.finditer(article_text))
    if not matches:
        return []
    items: List[tuple[str, str]] = []
    for idx, m in enumerate(matches):
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(article_text)
        clause_no = m.group(1)
        clause_body = article_text[start:end].strip()
        full_text = f"{article_header}\n{clause_body}".strip()
        items.append((clause_no, full_text))
    return items


def _split_points(clause_text: str) -> List[tuple[str, str]]:
    matches = list(POINT_RE.finditer(clause_text))
    if not matches:
        return []
    items: List[tuple[str, str]] = []
    for idx, m in enumerate(matches):
        start = m.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(clause_text)
        point_label = m.group(1)
        point_body = clause_text[start:end].strip()
        items.append((point_label, point_body))
    return items


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for DOCX ingestion") from exc
    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table_idx, table in enumerate(doc.tables, start=1):
        parts.append(f"[TABLE {table_idx}]")
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _guess_title(text: str, path: Path) -> str:
    for line in text.splitlines()[:20]:
        line = line.strip()
        if len(line) >= 12 and not line.lower().startswith(("cộng hòa", "độc lập")):
            return line[:180]
    return path.stem


def _guess_legal_type(text: str, path: Path) -> str:
    sample = text[:2000].lower() + " " + path.name.lower()
    if "nghị định" in sample:
        return "nghị định"
    if "thông tư" in sample:
        return "thông tư"
    if "luật" in sample:
        return "luật"
    return "unknown"


def _last_match(pattern: re.Pattern[str], text: str) -> str | None:
    matches = list(pattern.finditer(text))
    if not matches:
        return None
    return normalize_text(matches[-1].group(1))


def _first_line(text: str) -> str:
    for line in text.splitlines():
        line = normalize_text(line)
        if line:
            return line
    return ""
